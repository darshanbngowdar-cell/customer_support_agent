from collections.abc import Iterable
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from support_agent.domain.documents import SourceDocument


class UnsupportedDocumentTypeError(ValueError):
    """Raised when a file type is not supported by the knowledge base loader."""


class DocumentLoader:
    """Loads PDF, Markdown, TXT, and DOCX files while preserving useful metadata."""

    SUPPORTED_EXTENSIONS = {".pdf", ".md", ".markdown", ".txt", ".docx"}

    def load_directory(self, directory: Path) -> list[SourceDocument]:
        files = [
            path
            for path in directory.rglob("*")
            if path.is_file() and path.suffix.lower() in self.SUPPORTED_EXTENSIONS
        ]
        documents: list[SourceDocument] = []
        for file_path in sorted(files):
            documents.extend(self.load_file(file_path))
        return documents

    def load_file(self, file_path: Path) -> list[SourceDocument]:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._load_pdf(file_path)
        if suffix in {".md", ".markdown"}:
            return self._load_markdown(file_path)
        if suffix == ".txt":
            return self._load_txt(file_path)
        if suffix == ".docx":
            return self._load_docx(file_path)
        raise UnsupportedDocumentTypeError(f"Unsupported document type: {file_path.suffix}")

    def _base_metadata(self, file_path: Path, file_type: str) -> dict[str, str]:
        return {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": file_type,
        }

    def _load_pdf(self, file_path: Path) -> list[SourceDocument]:
        reader = PdfReader(str(file_path))
        documents: list[SourceDocument] = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                metadata = self._base_metadata(file_path, "pdf")
                metadata["page_number"] = page_index
                documents.append(SourceDocument(text=text, metadata=metadata))
        return documents

    def _load_markdown(self, file_path: Path) -> list[SourceDocument]:
        text = file_path.read_text(encoding="utf-8")
        return [
            SourceDocument(
                text=text,
                metadata={
                    **self._base_metadata(file_path, "markdown"),
                    "section_name": self._first_markdown_heading(text),
                },
            )
        ]

    def _load_txt(self, file_path: Path) -> list[SourceDocument]:
        text = file_path.read_text(encoding="utf-8")
        return [
            SourceDocument(
                text=text,
                metadata={
                    **self._base_metadata(file_path, "txt"),
                    "section_name": file_path.stem,
                },
            )
        ]

    def _load_docx(self, file_path: Path) -> list[SourceDocument]:
        doc = DocxDocument(str(file_path))
        sections: list[SourceDocument] = []
        current_heading = file_path.stem
        current_lines: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                sections.extend(self._docx_section(file_path, current_heading, current_lines))
                current_heading = text
                current_lines = []
            else:
                current_lines.append(text)

        sections.extend(self._docx_section(file_path, current_heading, current_lines))
        return sections

    def _docx_section(
        self,
        file_path: Path,
        heading: str,
        lines: Iterable[str],
    ) -> list[SourceDocument]:
        text = "\n".join(lines).strip()
        if not text:
            return []
        return [
            SourceDocument(
                text=text,
                metadata={
                    **self._base_metadata(file_path, "docx"),
                    "section_name": heading,
                },
            )
        ]

    @staticmethod
    def _first_markdown_heading(text: str) -> str:
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                return stripped.lstrip("#").strip()
        return "Document"
